from fastapi import FastAPI, APIRouter, HTTPException, Depends, BackgroundTasks, Response, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Literal
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from mailmerge import MailMerge
import random
import shutil
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'piperocket-secret-key-2025')
JWT_ALGORITHM = 'HS256'

app = FastAPI()
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

# ============= MODELS =============

class Organization(BaseModel):
    model_config = ConfigDict(extra="ignore")
    org_id: str = Field(default_factory=lambda: f"org_{uuid.uuid4().hex[:8]}")
    org_name: str
    logo: Optional[str] = None
    admin_name: str
    admin_email: EmailStr
    admin_mobile: str
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class OrganizationSignup(BaseModel):
    org_name: str
    admin_name: str
    admin_email: EmailStr
    admin_password: str
    admin_mobile: str

class UserRole(BaseModel):
    role: Literal['Admin', 'Director', 'Staff']

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    org_id: str
    name: str
    email: EmailStr
    mobile: Optional[str] = None
    role: Literal['Admin', 'Director', 'Staff']
    status: Literal['Active', 'Invited'] = 'Active'
    password_hash: Optional[str] = None
    otp_verified: bool = False
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class UserCreate(BaseModel):
    name: str
    email: EmailStr
    mobile: Optional[str] = None
    role: Literal['Admin', 'Director', 'Staff']
    password: str

class LoginRequest(BaseModel):
    org_id: str
    email: EmailStr
    password: str

class OTPVerifyRequest(BaseModel):
    email: EmailStr
    otp: str

class Client(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"client_{uuid.uuid4().hex[:8]}")
    org_id: str
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    end_date: str = ""
    currency_preference: Literal['USD', 'INR'] = 'INR'
    service: str  # Dynamic service from Service table
    amount_inr: float
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    authorised_signatory: str
    signatory_designation: str
    gst: str
    poc_name: str
    poc_email: EmailStr
    poc_designation: str
    poc_mobile: str
    approver_user_id: str
    sign_status: Literal['Signed', 'Not signed'] = 'Not signed'
    client_status: Literal['Active', 'Churned'] = 'Active'
    agreement_status: Literal['Live', 'Expired'] = 'Live'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ClientCreate(BaseModel):
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    currency_preference: Literal['USD', 'INR'] = 'INR'
    service: str  # Dynamic service from Service table
    amount_inr: float
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    authorised_signatory: str
    signatory_designation: str
    gst: str
    poc_name: str
    poc_email: EmailStr
    poc_designation: str
    poc_mobile: str
    approver_user_id: str

class Contractor(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"contractor_{uuid.uuid4().hex[:8]}")
    org_id: str
    name: str
    doj: str
    start_date: str
    tenure_months: int
    end_date: str = ""
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    pan: str
    aadhar: str
    mobile: str
    personal_email: EmailStr
    bank_name: str
    account_holder: str
    account_no: str
    ifsc: str
    address_1: str
    pincode: str
    city: str
    address_2: Optional[str] = None
    department: str  # Dynamic department from Service table
    projects: List[str] = Field(default_factory=list)
    monthly_retainer_inr: float
    designation: str
    approver_user_id: str
    sign_status: Literal['Signed', 'Not signed'] = 'Not signed'
    status: Literal['Active', 'Terminated'] = 'Active'
    agreement_status: Literal['Live', 'Expired'] = 'Live'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ContractorCreate(BaseModel):
    name: str
    doj: str
    start_date: str
    tenure_months: int
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    pan: str
    aadhar: str
    mobile: str
    personal_email: EmailStr
    bank_name: str
    account_holder: str
    account_no: str
    ifsc: str
    address_1: str
    pincode: str
    city: str
    address_2: Optional[str] = None
    department: str  # Dynamic department from Service table
    projects: List[str] = Field(default_factory=list)
    monthly_retainer_inr: float
    designation: str
    approver_user_id: str

class Employee(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"emp_{uuid.uuid4().hex[:8]}")
    org_id: str
    doj: str
    work_email: EmailStr
    emp_id: str
    first_name: str
    last_name: str
    father_name: str
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    mobile: str
    personal_email: EmailStr
    pan: str
    aadhar: str
    uan: str
    pf_account_no: str
    bank_name: str
    account_no: str
    ifsc: str
    branch: str
    address: str
    pincode: str
    city: str
    monthly_gross_inr: float
    department: str  # Dynamic department from Service table
    projects: List[str] = Field(default_factory=list)
    approver_user_id: str
    status: Literal['Active', 'Terminated'] = 'Active'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class EmployeeCreate(BaseModel):
    doj: str
    work_email: EmailStr
    emp_id: str
    first_name: str
    last_name: str
    father_name: str
    dob: str
    gender: Literal['Male', 'Female', 'Other'] = 'Male'
    mobile: str
    personal_email: EmailStr
    pan: str
    aadhar: str
    uan: str
    pf_account_no: str
    bank_name: str
    account_no: str
    ifsc: str
    branch: str
    address: str
    pincode: str
    city: str
    monthly_gross_inr: float
    department: str  # Dynamic department from Service table
    projects: List[str] = Field(default_factory=list)
    approver_user_id: str

class Approval(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"appr_{uuid.uuid4().hex[:8]}")
    item_type: Literal['client', 'contractor', 'employee']
    item_id: str
    requested_by: str
    status: Literal['Requested', 'Approved', 'Rejected', 'Hold'] = 'Requested'
    approved_by: Optional[str] = None
    approved_at: Optional[str] = None
    notes: Optional[str] = None
    staff_remarks: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ApprovalAction(BaseModel):
    action: Literal['approve', 'reject', 'hold']
    notes: Optional[str] = None

class ApprovalRequest(BaseModel):
    staff_remarks: Optional[str] = None

class SLAGenerateRequest(BaseModel):
    client_name: str
    address: str
    start_date: str
    tenure_months: int
    currency_preference: Literal['USD', 'INR']
    service: Literal['PPC', 'SEO', 'Content', 'Backlink']
    amount_ppc: Optional[float] = None
    amount_seo: Optional[float] = None
    amount: Optional[float] = None
    authorised_signatory: str
    designation: str

class NDAGenerateRequest(BaseModel):
    client_name: str
    address: str
    start_date: str
    authorised_signatory: str
    designation: str

class ICAGenerateRequest(BaseModel):
    contractor_name: str
    address: str
    start_date: str
    tenure_months: int
    amount_inr: float
    designation: str

class OfferLetterGenerateRequest(BaseModel):
    employee_name: str
    date: str
    gross_salary_lpa: float
    sign_before_date: str
    position: str
    department: str

class Asset(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"asset_{uuid.uuid4().hex[:8]}")
    org_id: str
    asset_type: str
    model: str
    serial_number: str
    purchase_date: str
    vendor: str
    value_ex_gst: float
    warranty_period_months: int
    alloted_to: str
    email: EmailStr
    department: str  # Dynamic department from Service table
    warranty_status: str = "Active"
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class AssetCreate(BaseModel):
    asset_type: str
    model: str
    serial_number: str
    purchase_date: str
    vendor: str
    value_ex_gst: float
    warranty_period_months: int
    alloted_to: str
    email: EmailStr
    department: str  # Dynamic department from Service table


# ============= SERVICE/DEPARTMENT MODELS =============
class Service(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"service_{uuid.uuid4().hex[:8]}")
    org_id: str
    name: str
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ServiceCreate(BaseModel):
    name: str

# ============= CLIENT ONBOARDING MODELS =============

class ClientOnboarding(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"onboard_{uuid.uuid4().hex[:8]}")
    org_id: str
    client_name: str
    poc_name: str
    poc_email: EmailStr
    services: List[str]  # Multiple: PPC, SEO, Backlink, Content
    currency: Literal['USD', 'INR']
    pricing: float
    approver_user_id: str
    proposal_status: Literal['Sent', 'Approved', 'Rejected', 'In Negotiation'] = 'Sent'
    onboarding_status: Literal['Onboarded', 'WIP', 'Not Onboarded'] = 'Not Onboarded'
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class ClientOnboardingCreate(BaseModel):
    client_name: str
    poc_name: str
    poc_email: EmailStr
    services: List[str]
    currency: Literal['USD', 'INR']
    pricing: float
    approver_user_id: str

# ============= CONSUMABLES MODELS =============

class StockAvailability(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"stock_{uuid.uuid4().hex[:8]}")
    org_id: str
    product_name: str
    vendor_name: str
    stock_available: int
    notes: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class StockTransaction(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: f"txn_{uuid.uuid4().hex[:8]}")
    org_id: str
    type: Literal['Stock In', 'Stock Out']
    product_name: str
    vendor_name_or_issued_to: str
    invoice_number: Optional[str] = None
    email: EmailStr
    date: str
    quantity: int
    price: Optional[float] = None  # For Stock In
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class StockInCreate(BaseModel):
    product_name: str
    quantity: int
    price: float
    vendor_name: str
    email: EmailStr
    invoice_number: str
    date: str

class StockOutCreate(BaseModel):
    product_name: str
    quantity: int
    issued_to: str
    email: EmailStr
    date: str

# ============= HELPER FUNCTIONS =============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str, org_id: str) -> str:
    payload = {
        'user_id': user_id,
        'email': email,
        'role': role,
        'org_id': org_id,
        'exp': datetime.now(timezone.utc) + timedelta(hours=24)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

def calculate_end_date(start_date: str, tenure_months: int) -> str:
    from dateutil.relativedelta import relativedelta
    start = datetime.fromisoformat(start_date)
    end = start + relativedelta(months=tenure_months)
    return end.isoformat()[:10]

def check_agreement_status(end_date: str) -> str:
    end = datetime.fromisoformat(end_date)
    # Remove timezone info for comparison
    today = datetime.now().date()
    end_date_only = end.date() if hasattr(end, 'date') else end
    return 'Live' if today <= end_date_only else 'Expired'

# ============= AUTH ROUTES =============

@api_router.post("/auth/signup")
async def signup(request: OrganizationSignup):
    # Check if email already exists
    existing_user = await db.users.find_one({"email": request.admin_email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create organization
    org = Organization(
        org_name=request.org_name,
        admin_name=request.admin_name,
        admin_email=request.admin_email,
        admin_mobile=request.admin_mobile
    )
    org_dict = org.model_dump()
    await db.organizations.insert_one(org_dict)
    
    # Create admin user
    user = User(
        org_id=org.org_id,
        name=request.admin_name,
        email=request.admin_email,
        mobile=request.admin_mobile,
        role='Admin',
        password_hash=hash_password(request.admin_password),
        otp_verified=False
    )
    user_dict = user.model_dump()
    await db.users.insert_one(user_dict)
    
    return {
        "message": "Organization created successfully",
        "org_id": org.org_id,
        "org_name": org.org_name,
        "admin_email": request.admin_email,
        "instructions": "Please use your Org ID, email, and password to login"
    }

@api_router.post("/auth/upload-logo")
async def upload_logo(org_id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    # Verify user belongs to this org and is admin
    if current_user.get('org_id') != org_id or current_user.get('role') != 'Admin':
        raise HTTPException(status_code=403, detail="Only organization admin can upload logo")
    
    # Create uploads directory if it doesn't exist
    uploads_dir = ROOT_DIR / "uploads" / "logos"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    
    # Save file
    file_extension = file.filename.split('.')[-1] if '.' in file.filename else 'png'
    filename = f"{org_id}.{file_extension}"
    file_path = uploads_dir / filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Update organization with logo path
    logo_url = f"/uploads/logos/{filename}"
    await db.organizations.update_one(
        {"org_id": org_id},
        {"$set": {"logo": logo_url}}
    )
    
    return {"message": "Logo uploaded successfully", "logo_url": logo_url}

@api_router.post("/auth/login")
async def login(request: LoginRequest):
    # Verify organization exists
    org = await db.organizations.find_one({"org_id": request.org_id})
    if not org:
        raise HTTPException(status_code=404, detail="Organization not found")
    
    # Find user with org_id and email
    user = await db.users.find_one({"org_id": request.org_id, "email": request.email})
    if not user or not verify_password(request.password, user.get('password_hash', '')):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Generate OTP (for MVP, we'll use a simple 6-digit code)
    otp = str(random.randint(100000, 999999))
    await db.otps.update_one(
        {"email": request.email, "org_id": request.org_id},
        {"$set": {"otp": otp, "created_at": datetime.now(timezone.utc).isoformat()}},
        upsert=True
    )
    
    return {
        "message": "OTP sent to email",
        "org_id": request.org_id,
        "email": request.email,
        "otp": otp,  # For MVP, returning OTP (in production, send via email)
        "requires_verification": not user.get('otp_verified', False)
    }

@api_router.post("/auth/verify-otp")
async def verify_otp(request: OTPVerifyRequest):
    otp_record = await db.otps.find_one({"email": request.email})
    if not otp_record or otp_record['otp'] != request.otp:
        raise HTTPException(status_code=400, detail="Invalid OTP")
    
    user = await db.users.find_one({"email": request.email})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Get organization info
    org = await db.organizations.find_one({"org_id": user['org_id']})
    
    # Mark user as verified
    await db.users.update_one(
        {"email": request.email},
        {"$set": {"otp_verified": True}}
    )
    
    # Generate JWT token with org_id
    token = create_token(user['id'], user['email'], user['role'], user['org_id'])
    
    return {
        "token": token,
        "org_id": user['org_id'],
        "org_name": org.get('org_name') if org else None,
        "org_logo": org.get('logo') if org else None,
        "user": {
            "id": user['id'],
            "name": user['name'],
            "email": user['email'],
            "role": user['role']
        }
    }

@api_router.get("/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    user = await db.users.find_one({"id": current_user['user_id']}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user

# ============= USER ROUTES =============

@api_router.get("/users", response_model=List[User])
async def get_users(current_user: dict = Depends(get_current_user)):
    # Only return users from the same organization
    users = await db.users.find({"org_id": current_user['org_id']}, {"_id": 0, "password_hash": 0}).to_list(1000)
    return users

@api_router.post("/users", response_model=User)
async def create_user(user_data: UserCreate, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can create users")
    
    # Check if user exists in this org
    existing = await db.users.find_one({"email": user_data.email, "org_id": current_user['org_id']})
    if existing:
        raise HTTPException(status_code=400, detail="User already exists in this organization")
    
    user = User(
        org_id=current_user['org_id'],  # Add org_id from current user
        name=user_data.name,
        email=user_data.email,
        mobile=user_data.mobile,
        role=user_data.role,
        password_hash=hash_password(user_data.password),
        status='Active'
    )
    
    doc = user.model_dump()
    await db.users.insert_one(doc)
    return user

@api_router.patch("/users/{user_id}")
async def update_user(user_id: str, role: str = None, status: str = None, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can update users")
    
    # Check if user being updated is admin
    user_to_update = await db.users.find_one({"id": user_id})
    if user_to_update and user_to_update.get('role') == 'Admin' and role and role != 'Admin':
        raise HTTPException(status_code=403, detail="Cannot change Admin role")
    
    update_data = {}
    if role:
        update_data['role'] = role
    if status:
        update_data['status'] = status
    
    await db.users.update_one({"id": user_id}, {"$set": update_data})
    return {"message": "User updated successfully"}

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can delete users")
    
    # Check if user being deleted is in same org
    user_to_delete = await db.users.find_one({"id": user_id, "org_id": current_user['org_id']})
    if not user_to_delete:
        raise HTTPException(status_code=404, detail="User not found in your organization")
    
    if user_to_delete.get('role') == 'Admin':
        raise HTTPException(status_code=403, detail="Cannot delete Admin user")
    
    result = await db.users.delete_one({"id": user_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="User not found")
    
    return {"message": "User deleted successfully"}


# ============= CLIENT ROUTES =============

@api_router.get("/clients", response_model=List[Client])
async def get_clients(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {"org_id": current_user['org_id']}  # Filter by org_id
    if filter_status:
        query['client_status'] = filter_status
    if filter_department:
        query['service'] = filter_department
    
    clients = await db.clients.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        clients.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return clients

@api_router.get("/clients/active-by-department")
async def get_active_clients_by_department(
    department: str = None,
    current_user: dict = Depends(get_current_user)
):
    """Get active clients filtered by service/department for project assignment"""
    query = {"client_status": "Active", "org_id": current_user['org_id']}  # Filter by org
    if department:
        query['service'] = department
    
    clients = await db.clients.find(query, {"_id": 0, "id": 1, "client_name": 1, "service": 1}).to_list(1000)
    return clients


@api_router.post("/clients", response_model=Client)
async def create_client(client_data: ClientCreate, current_user: dict = Depends(get_current_user)):
    client = Client(**client_data.model_dump(), org_id=current_user['org_id'])  # Add org_id
    client.end_date = calculate_end_date(client.start_date, client.tenure_months)
    client.agreement_status = check_agreement_status(client.end_date)
    
    doc = client.model_dump()
    await db.clients.insert_one(doc)
    return client

@api_router.patch("/clients/{client_id}")
async def update_client(client_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Verify client belongs to this org
    client = await db.clients.find_one({"id": client_id, "org_id": current_user['org_id']})
    if not client:
        raise HTTPException(status_code=404, detail="Client not found in your organization")
    
    # Recalculate end_date and agreement_status if start_date or tenure_months changed
    if 'start_date' in update_data or 'tenure_months' in update_data:
        start_date = update_data.get('start_date', client.get('start_date'))
        tenure_months = update_data.get('tenure_months', client.get('tenure_months'))
        
        if start_date and tenure_months:
            end_date = calculate_end_date(start_date, tenure_months)
            agreement_status = check_agreement_status(end_date)
            update_data['end_date'] = end_date
            update_data['agreement_status'] = agreement_status
    
    await db.clients.update_one({"id": client_id, "org_id": current_user['org_id']}, {"$set": update_data})
    return {"message": "Client updated successfully"}

@api_router.delete("/clients/{client_id}")
async def delete_client(client_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a client - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete clients")
    
    result = await db.clients.delete_one({"id": client_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Client not found in your organization")
    
    return {"message": "Client deleted successfully"}

@api_router.post("/clients/generate-sla")
async def generate_sla(request: SLAGenerateRequest):
    try:
        # Try template-based generation first
        if request.service == 'PPC':
            template_path = ROOT_DIR / 'templates' / 'SLA_PPC.docx'
        elif request.service == 'SEO':
            template_path = ROOT_DIR / 'templates' / 'SLA_SEO.docx'
        else:
            template_path = ROOT_DIR / 'templates' / 'SLA_PPC.docx'
        
        if template_path.exists():
            try:
                output_path = f"/tmp/SLA_{request.client_name.replace(' ', '_')}_{uuid.uuid4().hex[:6]}.docx"
                shutil.copy(template_path, output_path)
                
                merge_data = {
                    'client_name': request.client_name,
                    'address': request.address,
                    'start_date': request.start_date,
                    'tenure_months': str(request.tenure_months),
                    'service': request.service,
                    'currency': request.currency_preference,
                    'authorised_signatory': request.authorised_signatory,
                    'designation': request.designation,
                }
                
                if request.service == 'Both':
                    merge_data['amount_ppc'] = str(request.amount_ppc) if request.amount_ppc else '0'
                    merge_data['amount_seo'] = str(request.amount_seo) if request.amount_seo else '0'
                    merge_data['amount'] = str((request.amount_ppc or 0) + (request.amount_seo or 0))
                else:
                    merge_data['amount'] = str(request.amount) if request.amount else '0'
                
                document = MailMerge(output_path)
                document.merge(**merge_data)
                document.write(output_path)
                
                with open(output_path, 'rb') as f:
                    content = f.read()
                
                return Response(
                    content=content,
                    media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={'Content-Disposition': f'attachment; filename="SLA_{request.client_name.replace(" ", "_")}.docx"'}
                )
            except Exception as template_error:
                logger.error(f"Template merge error: {str(template_error)}")
                # Fall through to simple generation
        
    except Exception as e:
        logger.error(f"SLA generation error: {str(e)}")
    
    # Fallback: Generate simple document
    doc = Document()
    
    # Add header with logo placeholder
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('SERVICE LEVEL AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1  # Center
    
    doc.add_paragraph()  # Spacing
    
    # Agreement details
    doc.add_paragraph(f"This Service Level Agreement (\"SLA\") is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Client Name: {request.client_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph(f"Service Type: {request.service}")
    
    if request.service == 'Both':
        doc.add_paragraph(f"PPC Service Fee: {request.currency_preference} {request.amount_ppc:,.2f}")
        doc.add_paragraph(f"SEO Service Fee: {request.currency_preference} {request.amount_seo:,.2f}")
        doc.add_paragraph(f"Total Monthly Fee: {request.currency_preference} {(request.amount_ppc + request.amount_seo):,.2f}")
    else:
        doc.add_paragraph(f"Monthly Service Fee: {request.currency_preference} {request.amount:,.2f}")
    
    doc.add_paragraph(f"Contract Period: {request.tenure_months} months")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("For and on behalf of the Client:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.authorised_signatory}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="SLA_{request.client_name.replace(" ", "_")}.docx"'}
    )

@api_router.post("/clients/generate-nda")
async def generate_nda(request: NDAGenerateRequest):
    # Generate simple NDA document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('NON-DISCLOSURE AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    
    doc.add_paragraph(f"This Non-Disclosure Agreement (\"NDA\") is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Client Name: {request.client_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph()
    
    doc.add_paragraph("This agreement governs the disclosure of confidential information between the parties.")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("For and on behalf of the Client:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.authorised_signatory}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="NDA_{request.client_name.replace(" ", "_")}.docx"'}
    )

# ============= CONTRACTOR ROUTES =============

@api_router.get("/contractors", response_model=List[Contractor])
async def get_contractors(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {"org_id": current_user['org_id']}  # Filter by org_id
    if filter_status:
        query['status'] = filter_status
    if filter_department:
        query['department'] = filter_department
    
    contractors = await db.contractors.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        contractors.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return contractors

@api_router.post("/contractors", response_model=Contractor)
async def create_contractor(contractor_data: ContractorCreate, current_user: dict = Depends(get_current_user)):
    contractor = Contractor(**contractor_data.model_dump(), org_id=current_user['org_id'])  # Add org_id
    contractor.end_date = calculate_end_date(contractor.start_date, contractor.tenure_months)
    contractor.agreement_status = check_agreement_status(contractor.end_date)
    
    doc = contractor.model_dump()
    await db.contractors.insert_one(doc)
    return contractor

@api_router.patch("/contractors/{contractor_id}")
async def update_contractor(contractor_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Verify contractor belongs to this org
    contractor = await db.contractors.find_one({"id": contractor_id, "org_id": current_user['org_id']})
    if not contractor:
        raise HTTPException(status_code=404, detail="Contractor not found in your organization")
    
    # Recalculate end_date and agreement_status if start_date or tenure_months changed
    if 'start_date' in update_data or 'tenure_months' in update_data:
        start_date = update_data.get('start_date', contractor.get('start_date'))
        tenure_months = update_data.get('tenure_months', contractor.get('tenure_months'))
        
        if start_date and tenure_months:
            end_date = calculate_end_date(start_date, tenure_months)
            agreement_status = check_agreement_status(end_date)
            update_data['end_date'] = end_date
            update_data['agreement_status'] = agreement_status
    
    await db.contractors.update_one({"id": contractor_id, "org_id": current_user['org_id']}, {"$set": update_data})
    return {"message": "Contractor updated successfully"}

@api_router.delete("/contractors/{contractor_id}")
async def delete_contractor(contractor_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a contractor - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete contractors")
    
    result = await db.contractors.delete_one({"id": contractor_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Contractor not found in your organization")
    
    return {"message": "Contractor deleted successfully"}

@api_router.post("/contractors/generate-ica")
async def generate_ica(request: ICAGenerateRequest):
    # Generate simple ICA document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('INDEPENDENT CONTRACTOR AGREEMENT')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    
    doc.add_paragraph(f"This Independent Contractor Agreement is entered into on {request.start_date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Contractor Name: {request.contractor_name}")
    doc.add_paragraph(f"Address: {request.address}")
    doc.add_paragraph(f"Designation: {request.designation}")
    doc.add_paragraph(f"Monthly Retainer: INR {request.amount_inr:,.2f}")
    doc.add_paragraph(f"Contract Period: {request.tenure_months} months")
    doc.add_paragraph()
    
    doc.add_paragraph("Terms and Conditions:")
    doc.add_paragraph("1. The Contractor agrees to provide services as per the scope of work.")
    doc.add_paragraph("2. Payment shall be made on a monthly basis.")
    doc.add_paragraph("3. This agreement can be terminated by either party with 30 days notice.")
    doc.add_paragraph()
    
    # Signature section
    doc.add_paragraph("Contractor Signature:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.contractor_name}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="ICA_{request.contractor_name.replace(" ", "_")}.docx"'}
    )

# ============= EMPLOYEE ROUTES =============

@api_router.get("/employees", response_model=List[Employee])
async def get_employees(
    current_user: dict = Depends(get_current_user),
    sort_by: str = None,
    sort_order: str = 'asc',
    filter_status: str = None,
    filter_department: str = None
):
    query = {"org_id": current_user['org_id']}  # Filter by org_id
    if filter_status:
        query['status'] = filter_status
    if filter_department:
        query['department'] = filter_department
    
    employees = await db.employees.find(query, {"_id": 0}).to_list(1000)
    
    # Sorting
    if sort_by:
        reverse = sort_order == 'desc'
        employees.sort(key=lambda x: x.get(sort_by, ''), reverse=reverse)
    
    return employees

@api_router.post("/employees", response_model=Employee)
async def create_employee(employee_data: EmployeeCreate, current_user: dict = Depends(get_current_user)):
    employee = Employee(**employee_data.model_dump(), org_id=current_user['org_id'])  # Add org_id
    
    doc = employee.model_dump()
    await db.employees.insert_one(doc)
    return employee

@api_router.patch("/employees/{employee_id}")
async def update_employee(employee_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Verify employee belongs to this org
    employee = await db.employees.find_one({"id": employee_id, "org_id": current_user['org_id']})
    if not employee:
        raise HTTPException(status_code=404, detail="Employee not found in your organization")
    
    await db.employees.update_one({"id": employee_id, "org_id": current_user['org_id']}, {"$set": update_data})
    return {"message": "Employee updated successfully"}

@api_router.delete("/employees/{employee_id}")
async def delete_employee(employee_id: str, current_user: dict = Depends(get_current_user)):
    """Delete an employee - Admin and Director only"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete employees")
    
    result = await db.employees.delete_one({"id": employee_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Employee not found in your organization")
    
    return {"message": "Employee deleted successfully"}

@api_router.post("/employees/generate-offer")
async def generate_offer_letter(request: OfferLetterGenerateRequest):
    # Calculate CTC
    gross_annual = request.gross_salary_lpa * 100000
    ctc_annual = gross_annual + 21600
    monthly_ctc = ctc_annual / 12
    monthly_gross = gross_annual / 12
    
    # Generate offer letter document
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "PIPEROCKET"
    header_para.style.font.size = Pt(16)
    header_para.style.font.bold = True
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('OFFER LETTER')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = 1
    
    doc.add_paragraph()
    doc.add_paragraph(f"Date: {request.date}")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Dear {request.employee_name},")
    doc.add_paragraph()
    
    doc.add_paragraph(f"We are pleased to offer you the position of {request.position} in the {request.department} department.")
    doc.add_paragraph()
    
    doc.add_paragraph("Compensation Details:")
    doc.add_paragraph(f"• Gross Annual Salary: INR {gross_annual:,.2f}")
    doc.add_paragraph(f"• Cost to Company (Annual): INR {ctc_annual:,.2f}")
    doc.add_paragraph(f"• Monthly CTC: INR {monthly_ctc:,.2f}")
    doc.add_paragraph(f"• Monthly Gross: INR {monthly_gross:,.2f}")
    doc.add_paragraph()
    
    # Salary breakdown table
    doc.add_paragraph("Monthly Salary Breakdown:")
    
    # Calculate components
    basic = monthly_gross * 0.50
    hra = monthly_gross * 0.30
    special = monthly_gross * 0.20
    
    doc.add_paragraph(f"• Basic Salary: INR {basic:,.2f}")
    doc.add_paragraph(f"• HRA: INR {hra:,.2f}")
    doc.add_paragraph(f"• Special Allowance: INR {special:,.2f}")
    doc.add_paragraph(f"• Employer PF Contribution: INR 1,800.00")
    doc.add_paragraph()
    
    doc.add_paragraph(f"Please sign and return this offer letter before {request.sign_before_date}.")
    doc.add_paragraph()
    
    doc.add_paragraph("We look forward to welcoming you to our team!")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Piperocket HR Team")
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph("Employee Acceptance:")
    doc.add_paragraph()
    doc.add_paragraph(f"Name: {request.employee_name}")
    doc.add_paragraph(f"Date: ________________")
    doc.add_paragraph(f"Signature: ________________")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    
    return Response(
        content=bio.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="Offer_{request.employee_name.replace(" ", "_")}.docx"'}
    )

# ============= APPROVAL ROUTES =============

@api_router.get("/approvals", response_model=List[Approval])
async def get_approvals(current_user: dict = Depends(get_current_user)):
    # Get approvals for items belonging to this org
    # For now, we'll need to check the actual items, but for simplicity filter approvals by users in org
    org_users = await db.users.find({"org_id": current_user['org_id']}, {"id": 1}).to_list(1000)
    org_user_ids = [u['id'] for u in org_users]
    approvals = await db.approvals.find({"requested_by": {"$in": org_user_ids}}, {"_id": 0}).to_list(1000)
    return approvals

@api_router.post("/approvals/{item_type}/{item_id}/request")
async def request_approval(item_type: str, item_id: str, request: ApprovalRequest, current_user: dict = Depends(get_current_user)):
    if current_user['role'] == 'Director':
        raise HTTPException(status_code=403, detail="Directors cannot request approval")
    
    approval = Approval(
        item_type=item_type,
        item_id=item_id,
        requested_by=current_user['user_id'],
        status='Requested',
        staff_remarks=request.staff_remarks
    )
    
    doc = approval.model_dump()
    await db.approvals.insert_one(doc)
    return approval

@api_router.post("/approvals/{approval_id}/action")
async def approval_action(approval_id: str, action: ApprovalAction, current_user: dict = Depends(get_current_user)):
    if current_user['role'] != 'Director':
        raise HTTPException(status_code=403, detail="Only Directors can approve/reject/hold")
    
    status_map = {'approve': 'Approved', 'reject': 'Rejected', 'hold': 'Hold'}
    status = status_map.get(action.action, 'Requested')
    
    await db.approvals.update_one(
        {"id": approval_id},
        {"$set": {
            "status": status,
            "approved_by": current_user['user_id'],
            "approved_at": datetime.now(timezone.utc).isoformat(),
            "notes": action.notes
        }}
    )
    
    return {"message": f"Approval {status.lower()} successfully"}

@api_router.delete("/approvals/reset")
async def reset_approvals(current_user: dict = Depends(get_current_user)):
    """Reset all approval records - accessible by Staff and Admin"""
    if current_user['role'] == 'Director':
        raise HTTPException(status_code=403, detail="Directors cannot reset approvals")
    
    # Only reset approvals for this org
    org_users = await db.users.find({"org_id": current_user['org_id']}, {"id": 1}).to_list(1000)
    org_user_ids = [u['id'] for u in org_users]
    result = await db.approvals.delete_many({"requested_by": {"$in": org_user_ids}})
    return {"message": f"Reset complete. Deleted {result.deleted_count} approval records"}

# ============= DASHBOARD ROUTES =============

@api_router.get("/dashboard/summary")
async def get_dashboard_summary(current_user: dict = Depends(get_current_user)):
    # Get alerts
    today = datetime.now(timezone.utc)
    thirty_days_later = today + timedelta(days=30)
    
    # Expiring agreements - get actual client names (filter by org_id)
    clients_all = await db.clients.find({"client_status": "Active", "org_id": current_user['org_id']}).to_list(1000)
    expiring_clients = []
    for client in clients_all:
        if client.get('end_date'):
            try:
                end_date = datetime.fromisoformat(client['end_date'])
                # Remove timezone for comparison
                if hasattr(end_date, 'date'):
                    end_date_only = end_date.date()
                else:
                    end_date_only = end_date
                
                if today.date() <= end_date_only <= thirty_days_later.date():
                    expiring_clients.append({
                        "name": client['client_name'],
                        "end_date": client['end_date'],
                        "service": client['service']
                    })
            except:
                pass
    
    # Upcoming birthdays - get employee and contractor names (filter by org_id)
    employees = await db.employees.find({"status": "Active", "org_id": current_user['org_id']}).to_list(1000)
    contractors = await db.contractors.find({"status": "Active", "org_id": current_user['org_id']}).to_list(1000)
    
    upcoming_birthdays = []
    
    for emp in employees:
        if emp.get('dob'):
            try:
                dob = datetime.fromisoformat(emp['dob'])
                # Calculate days until birthday this year
                this_year_bday = datetime(today.year, dob.month, dob.day)
                if this_year_bday < datetime.now():
                    # Birthday already passed this year, check next year
                    this_year_bday = datetime(today.year + 1, dob.month, dob.day)
                
                days_until = (this_year_bday - datetime.now()).days
                if 0 <= days_until <= 15:
                    upcoming_birthdays.append({
                        "name": f"{emp['first_name']} {emp['last_name']}",
                        "date": emp['dob'],
                        "type": "Employee",
                        "department": emp.get('department', '')
                    })
            except:
                pass
    
    for con in contractors:
        if con.get('dob'):
            try:
                dob = datetime.fromisoformat(con['dob'])
                this_year_bday = datetime(today.year, dob.month, dob.day)
                if this_year_bday < datetime.now():
                    this_year_bday = datetime(today.year + 1, dob.month, dob.day)
                
                days_until = (this_year_bday - datetime.now()).days
                if 0 <= days_until <= 15:
                    upcoming_birthdays.append({
                        "name": con['name'],
                        "date": con['dob'],
                        "type": "Contractor",
                        "department": con.get('department', '')
                    })
            except:
                pass
    
    # Sort birthdays by date
    upcoming_birthdays.sort(key=lambda x: datetime.fromisoformat(x['date']))
    
    # Expired agreements
    expired_clients = []
    for client in clients_all:
        if client.get('end_date'):
            try:
                end_date = datetime.fromisoformat(client['end_date'])
                if hasattr(end_date, 'date'):
                    end_date_only = end_date.date()
                else:
                    end_date_only = end_date
                
                if end_date_only < today.date():
                    expired_clients.append({
                        "name": client['client_name'],
                        "end_date": client['end_date'],
                        "service": client['service']
                    })
            except:
                pass
    
    # Revenue metrics (filter by org_id)
    clients = await db.clients.find({"client_status": "Active", "org_id": current_user['org_id']}).to_list(1000)
    revenue_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_clients = [c for c in clients if c.get('service') == dept]
        count = len(dept_clients)
        amount = sum(c.get('amount_inr', 0) for c in dept_clients)
        revenue_by_dept[dept] = {"count": count, "amount": amount}
    
    # Employee metrics
    employee_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_employees = [e for e in employees if e.get('department') == dept]
        count = len(dept_employees)
        cost = sum(e.get('monthly_gross_inr', 0) for e in dept_employees)
        employee_by_dept[dept] = {"count": count, "cost": cost}
    
    # Contractor metrics
    contractor_by_dept = {}
    for dept in ['PPC', 'SEO', 'Content', 'Backlink', 'Business Development', 'Others']:
        dept_contractors = [c for c in contractors if c.get('department') == dept]
        count = len(dept_contractors)
        cost = sum(c.get('monthly_retainer_inr', 0) for c in dept_contractors)
        contractor_by_dept[dept] = {"count": count, "cost": cost}
    
    return {
        "alerts": {
            "expiring_agreements": expiring_clients,
            "expired_agreements": expired_clients,
            "upcoming_birthdays": upcoming_birthdays
        },
        "revenue": revenue_by_dept,
        "employees": employee_by_dept,
        "contractors": contractor_by_dept
    }

# ============= BULK EXPORT/IMPORT ROUTES =============

@api_router.get("/clients/export")
async def export_clients(current_user: dict = Depends(get_current_user)):
    """Export all clients to Excel"""
    clients = await db.clients.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    
    if not clients:
        raise HTTPException(status_code=404, detail="No clients to export")
    
    df = pd.DataFrame(clients)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Clients')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="clients_export.xlsx"'}
    )

@api_router.get("/contractors/export")
async def export_contractors(current_user: dict = Depends(get_current_user)):
    """Export all contractors to Excel"""
    contractors = await db.contractors.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    
    if not contractors:
        raise HTTPException(status_code=404, detail="No contractors to export")
    
    df = pd.DataFrame(contractors)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Contractors')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="contractors_export.xlsx"'}
    )

@api_router.get("/employees/export")
async def export_employees(current_user: dict = Depends(get_current_user)):
    """Export all employees to Excel"""
    employees = await db.employees.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    
    if not employees:
        raise HTTPException(status_code=404, detail="No employees to export")
    
    df = pd.DataFrame(employees)
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Employees')
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="employees_export.xlsx"'}
    )

@api_router.get("/clients/sample")
async def get_client_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    # Get a valid Director/Admin user ID for the sample
    approver = await db.users.find_one(
        {"org_id": current_user['org_id'], "role": {"$in": ["Admin", "Director"]}},
        {"id": 1}
    )
    approver_id = approver['id'] if approver else current_user['id']
    
    sample_data = {
        'client_name': ['ABC Corp', 'XYZ Ltd'],
        'address': ['123 Main St, New York', '456 Park Ave, Boston'],
        'start_date': ['2025-01-01', '2025-02-01'],
        'tenure_months': [12, 6],
        'currency_preference': ['INR', 'INR'],
        'service': ['PPC', 'SEO'],
        'amount_inr': [50000.0, 75000.0],
        'authorised_signatory': ['John Doe', 'Mike Johnson'],
        'signatory_designation': ['CEO', 'Director'],
        'gst': ['GST123456', 'GST789012'],
        'poc_name': ['Jane Smith', 'Sarah Lee'],
        'poc_email': ['jane.smith@abccorp.com', 'sarah.lee@xyzltd.com'],
        'poc_designation': ['Manager', 'Lead'],
        'poc_mobile': ['9876543210', '9876543211'],
        'approver_user_id': [approver_id, approver_id]
    }
    
    df = pd.DataFrame(sample_data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Clients')
        
        workbook = writer.book
        worksheet = writer.sheets['Clients']
        
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="client_sample.xlsx"'}
    )

@api_router.get("/contractors/sample")
async def get_contractor_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    # Get a valid Director/Admin user ID for the sample
    approver = await db.users.find_one(
        {"org_id": current_user['org_id'], "role": {"$in": ["Admin", "Director"]}},
        {"id": 1}
    )
    approver_id = approver['id'] if approver else current_user['id']
    
    sample_data = {
        'name': ['John Contractor', 'Mary Freelancer'],
        'doj': ['2025-01-01', '2025-01-15'],
        'start_date': ['2025-01-01', '2025-01-15'],
        'tenure_months': [6, 12],
        'dob': ['1990-05-15', '1992-08-20'],
        'pan': ['ABCDE1234F', 'XYZAB5678C'],
        'aadhar': ['123456789012', '987654321098'],
        'mobile': ['9876543210', '9876543211'],
        'personal_email': ['john.contractor@email.com', 'mary.freelancer@email.com'],
        'bank_name': ['HDFC Bank', 'ICICI Bank'],
        'account_holder': ['John Contractor', 'Mary Freelancer'],
        'account_no': ['1234567890', '0987654321'],
        'ifsc': ['HDFC0001234', 'ICIC0005678'],
        'address_1': ['123 Street, Area', '456 Avenue, Sector'],
        'pincode': ['110001', '110002'],
        'city': ['Delhi', 'Mumbai'],
        'address_2': ['Near Market', 'Behind Mall'],
        'department': ['PPC', 'SEO'],
        'monthly_retainer_inr': [35000.0, 40000.0],
        'designation': ['Consultant', 'Specialist'],
        'approver_user_id': [approver_id, approver_id]
    }
    
    df = pd.DataFrame(sample_data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Contractors')
        
        workbook = writer.book
        worksheet = writer.sheets['Contractors']
        
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="contractor_sample.xlsx"'}
    )

@api_router.get("/employees/sample")
async def get_employee_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    # Get a valid Director/Admin user ID for the sample
    approver = await db.users.find_one(
        {"org_id": current_user['org_id'], "role": {"$in": ["Admin", "Director"]}},
        {"id": 1}
    )
    approver_id = approver['id'] if approver else current_user['id']
    
    sample_data = {
        'doj': ['2025-01-15', '2025-02-01'],
        'work_email': ['john.doe@company.com', 'jane.smith@company.com'],
        'emp_id': ['EMP001', 'EMP002'],
        'first_name': ['John', 'Jane'],
        'last_name': ['Doe', 'Smith'],
        'father_name': ['James Doe', 'Robert Smith'],
        'dob': ['1995-03-20', '1993-07-15'],
        'mobile': ['9876543210', '9876543211'],
        'personal_email': ['john.personal@email.com', 'jane.personal@email.com'],
        'pan': ['ABCDE1234F', 'XYZAB5678C'],
        'aadhar': ['123456789012', '987654321098'],
        'uan': ['UAN123456', 'UAN789012'],
        'pf_account_no': ['PF123456', 'PF789012'],
        'bank_name': ['HDFC Bank', 'ICICI Bank'],
        'account_no': ['1234567890', '0987654321'],
        'ifsc': ['HDFC0001234', 'ICIC0005678'],
        'branch': ['Main Branch', 'City Branch'],
        'address': ['123 Street, Area', '456 Avenue, Sector'],
        'pincode': ['110001', '110002'],
        'city': ['Delhi', 'Mumbai'],
        'monthly_gross_inr': [60000.0, 75000.0],
        'department': ['PPC', 'SEO'],
        'approver_user_id': [approver_id, approver_id]
    }
    
    df = pd.DataFrame(sample_data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Employees')
        
        workbook = writer.book
        worksheet = writer.sheets['Employees']
        
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="employee_sample.xlsx"'}
    )

@api_router.get("/assets/sample")
async def get_asset_sample(current_user: dict = Depends(get_current_user)):
    """Download sample Excel template for bulk upload"""
    # Use pandas to create a proper DataFrame-based Excel file
    sample_data = {
        'asset_type': ['Laptop', 'Monitor', 'Keyboard'],
        'model': ['Dell XPS 15', 'LG 27inch 4K', 'Logitech MX Keys'],
        'serial_number': ['SN123456', 'SN789012', 'SN345678'],
        'purchase_date': ['2024-01-15', '2024-02-01', '2024-03-10'],
        'vendor': ['Dell India', 'LG Store', 'Amazon'],
        'value_ex_gst': [75000.0, 15000.0, 8500.0],
        'warranty_period_months': [12, 24, 12],
        'alloted_to': ['John Doe', 'Jane Smith', 'Bob Wilson'],
        'email': ['john.doe@example.com', 'jane.smith@example.com', 'bob.wilson@example.com'],
        'department': ['PPC', 'SEO', 'Content']
    }
    
    df = pd.DataFrame(sample_data)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Assets')
        
        # Format the header row
        workbook = writer.book
        worksheet = writer.sheets['Assets']
        
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="asset_sample.xlsx"'}
    )

@api_router.post("/clients/import")
async def import_clients(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import clients from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Handle start_date - could be datetime, date, or string
                start_date_val = row['start_date']
                if isinstance(start_date_val, (pd.Timestamp, datetime)):
                    start_date_str = start_date_val.strftime('%Y-%m-%d')
                else:
                    start_date_str = str(start_date_val)[:10]
                
                # Create client object
                client_data = ClientCreate(
                    client_name=str(row['client_name']).strip(),
                    address=str(row['address']).strip(),
                    start_date=start_date_str,
                    tenure_months=int(row['tenure_months']),
                    currency_preference=str(row.get('currency_preference', 'INR')).strip(),
                    service=str(row['service']).strip(),
                    amount_inr=float(row['amount_inr']),
                    authorised_signatory=str(row['authorised_signatory']).strip(),
                    signatory_designation=str(row['signatory_designation']).strip(),
                    gst=str(row['gst']).strip(),
                    poc_name=str(row['poc_name']).strip(),
                    poc_email=str(row['poc_email']).strip().lower(),
                    poc_designation=str(row['poc_designation']).strip(),
                    poc_mobile=str(row['poc_mobile']).strip(),
                    approver_user_id=str(row['approver_user_id']).strip()
                )
                
                client = Client(**client_data.model_dump(), org_id=current_user['org_id'])
                client.end_date = calculate_end_date(client.start_date, client.tenure_months)
                client.agreement_status = check_agreement_status(client.end_date)
                
                await db.clients.insert_one(client.model_dump())
                imported_count += 1
                
            except Exception as e:
                error_msg = f"Row {index + 2}: {str(e)}"
                errors.append(error_msg)
                print(f"Client import error: {error_msg}")
        
        return {
            "message": f"Import completed. {imported_count} clients imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/contractors/import")
async def import_contractors(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import contractors from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Handle date fields
                def parse_date(date_val):
                    if isinstance(date_val, (pd.Timestamp, datetime)):
                        return date_val.strftime('%Y-%m-%d')
                    return str(date_val)[:10]
                
                contractor_data = ContractorCreate(
                    name=str(row['name']).strip(),
                    doj=parse_date(row['doj']),
                    start_date=parse_date(row['start_date']),
                    tenure_months=int(row['tenure_months']),
                    dob=parse_date(row['dob']),
                    pan=str(row['pan']).strip().upper(),
                    aadhar=str(row['aadhar']).strip(),
                    mobile=str(row['mobile']).strip(),
                    personal_email=str(row['personal_email']).strip().lower(),
                    bank_name=str(row['bank_name']).strip(),
                    account_holder=str(row['account_holder']).strip(),
                    account_no=str(row['account_no']).strip(),
                    ifsc=str(row['ifsc']).strip().upper(),
                    address_1=str(row['address_1']).strip(),
                    pincode=str(row['pincode']).strip(),
                    city=str(row['city']).strip(),
                    address_2=str(row.get('address_2', '')).strip(),
                    department=str(row['department']).strip(),
                    monthly_retainer_inr=float(row['monthly_retainer_inr']),
                    designation=str(row['designation']).strip(),
                    approver_user_id=str(row['approver_user_id']).strip()
                )
                
                contractor = Contractor(**contractor_data.model_dump(), org_id=current_user['org_id'])
                contractor.end_date = calculate_end_date(contractor.start_date, contractor.tenure_months)
                contractor.agreement_status = check_agreement_status(contractor.end_date)
                
                await db.contractors.insert_one(contractor.model_dump())
                imported_count += 1
                
            except Exception as e:
                error_msg = f"Row {index + 2}: {str(e)}"
                errors.append(error_msg)
                print(f"Contractor import error: {error_msg}")
        
        return {
            "message": f"Import completed. {imported_count} contractors imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/employees/import")
async def import_employees(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import employees from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Handle date fields
                def parse_date(date_val):
                    if isinstance(date_val, (pd.Timestamp, datetime)):
                        return date_val.strftime('%Y-%m-%d')
                    return str(date_val)[:10]
                
                employee_data = EmployeeCreate(
                    doj=parse_date(row['doj']),
                    work_email=str(row['work_email']).strip().lower(),
                    emp_id=str(row['emp_id']).strip().upper(),
                    first_name=str(row['first_name']).strip(),
                    last_name=str(row['last_name']).strip(),
                    father_name=str(row['father_name']).strip(),
                    dob=parse_date(row['dob']),
                    mobile=str(row['mobile']).strip(),
                    personal_email=str(row['personal_email']).strip().lower(),
                    pan=str(row['pan']).strip().upper(),
                    aadhar=str(row['aadhar']).strip(),
                    uan=str(row['uan']).strip(),
                    pf_account_no=str(row['pf_account_no']).strip(),
                    bank_name=str(row['bank_name']).strip(),
                    account_no=str(row['account_no']).strip(),
                    ifsc=str(row['ifsc']).strip().upper(),
                    branch=str(row['branch']).strip(),
                    address=str(row['address']).strip(),
                    pincode=str(row['pincode']).strip(),
                    city=str(row['city']).strip(),
                    monthly_gross_inr=float(row['monthly_gross_inr']),
                    department=str(row['department']).strip(),
                    approver_user_id=str(row['approver_user_id']).strip()
                )
                
                employee = Employee(**employee_data.model_dump(), org_id=current_user['org_id'])
                await db.employees.insert_one(employee.model_dump())
                imported_count += 1
                
            except Exception as e:
                error_msg = f"Row {index + 2}: {str(e)}"
                errors.append(error_msg)
                print(f"Employee import error: {error_msg}")
        
        return {
            "message": f"Import completed. {imported_count} employees imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

@api_router.post("/assets/import")
async def import_assets(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """Bulk import assets from Excel"""
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can bulk upload")
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="Only Excel files (.xlsx, .xls) are supported")
    
    try:
        contents = await file.read()
        df = pd.read_excel(BytesIO(contents))
        
        imported_count = 0
        errors = []
        
        for index, row in df.iterrows():
            try:
                # Handle purchase_date - could be datetime, date, or string
                purchase_date_val = row['purchase_date']
                if pd.isna(purchase_date_val):
                    raise ValueError("purchase_date is required")
                
                # Convert to string date format
                if isinstance(purchase_date_val, (pd.Timestamp, datetime)):
                    purchase_date_str = purchase_date_val.strftime('%Y-%m-%d')
                else:
                    purchase_date_str = str(purchase_date_val)
                    if 'T' in purchase_date_str or ' ' in purchase_date_str:
                        purchase_date_str = purchase_date_str.split('T')[0].split(' ')[0]
                
                # Validate and clean data
                asset_data = AssetCreate(
                    asset_type=str(row['asset_type']).strip(),
                    model=str(row['model']).strip(),
                    serial_number=str(row['serial_number']).strip(),
                    purchase_date=purchase_date_str,
                    vendor=str(row['vendor']).strip(),
                    value_ex_gst=float(row['value_ex_gst']),
                    warranty_period_months=int(row['warranty_period_months']),
                    alloted_to=str(row['alloted_to']).strip(),
                    email=str(row['email']).strip().lower(),
                    department=str(row['department']).strip()
                )
                
                asset = Asset(**asset_data.model_dump(), org_id=current_user['org_id'])
                
                # Calculate warranty status - use simple date comparison
                purchase_date = datetime.strptime(asset.purchase_date, '%Y-%m-%d').date()
                warranty_end = purchase_date + timedelta(days=asset.warranty_period_months * 30)
                today = datetime.now().date()
                asset.warranty_status = 'Active' if today <= warranty_end else 'Expired'
                
                await db.assets.insert_one(asset.model_dump())
                imported_count += 1
                
            except Exception as e:
                error_msg = f"Row {index + 2}: {str(e)}"
                errors.append(error_msg)
                print(f"Import error: {error_msg}")  # Log to console for debugging
        
        return {
            "message": f"Import completed. {imported_count} assets imported successfully.",
            "imported": imported_count,
            "errors": errors if errors else None
        }
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")

# ============= ASSET TRACKER ROUTES =============

@api_router.get("/assets", response_model=List[Asset])
async def get_assets(
    current_user: dict = Depends(get_current_user),
    department: str = None
):
    query = {"org_id": current_user['org_id']}  # Filter by org_id
    if department:
        query['department'] = department
    
    assets = await db.assets.find(query, {"_id": 0}).to_list(1000)
    
    # Update warranty status
    for asset in assets:
        purchase_date = datetime.fromisoformat(asset['purchase_date'])
        warranty_end = purchase_date + timedelta(days=asset['warranty_period_months'] * 30)
        if datetime.now() > warranty_end:
            asset['warranty_status'] = 'Expired'
        else:
            asset['warranty_status'] = 'Active'
    
    return assets

@api_router.post("/assets", response_model=Asset)
async def create_asset(asset_data: AssetCreate, current_user: dict = Depends(get_current_user)):
    asset = Asset(**asset_data.model_dump(), org_id=current_user['org_id'])  # Add org_id
    
    # Calculate warranty status
    purchase_date = datetime.fromisoformat(asset.purchase_date)
    warranty_end = purchase_date + timedelta(days=asset.warranty_period_months * 30)
    asset.warranty_status = 'Active' if datetime.now() <= warranty_end else 'Expired'
    
    doc = asset.model_dump()
    await db.assets.insert_one(doc)
    return asset

@api_router.patch("/assets/{asset_id}")
async def update_asset(asset_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    # Verify asset belongs to this org
    asset = await db.assets.find_one({"id": asset_id, "org_id": current_user['org_id']})
    if not asset:
        raise HTTPException(status_code=404, detail="Asset not found in your organization")
    
    await db.assets.update_one({"id": asset_id, "org_id": current_user['org_id']}, {"$set": update_data})
    return {"message": "Asset updated successfully"}

@api_router.delete("/assets/{asset_id}")
async def delete_asset(asset_id: str, current_user: dict = Depends(get_current_user)):
    if current_user['role'] not in ['Admin', 'Director']:
        raise HTTPException(status_code=403, detail="Only Admin and Director can delete assets")
    
    result = await db.assets.delete_one({"id": asset_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Asset not found in your organization")
    
    return {"message": "Asset deleted successfully"}

@api_router.get("/assets/export")
async def export_assets(current_user: dict = Depends(get_current_user)):
    """Export all assets to Excel"""
    assets = await db.assets.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    
    if not assets:
        raise HTTPException(status_code=404, detail="No assets to export")
    
    # Create DataFrame and select only the columns needed for import
    df = pd.DataFrame(assets)
    
    # Select and reorder columns to match import format
    export_columns = ['asset_type', 'model', 'serial_number', 'purchase_date', 'vendor', 
                     'value_ex_gst', 'warranty_period_months', 'alloted_to', 'email', 'department']
    
    # Only include columns that exist
    df = df[[col for col in export_columns if col in df.columns]]
    
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Assets')
        
        # Format the header row
        workbook = writer.book
        worksheet = writer.sheets['Assets']
        
        for cell in worksheet[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="CCE5FF", end_color="CCE5FF", fill_type="solid")
    
    output.seek(0)
    
    return Response(
        content=output.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': 'attachment; filename="assets_export.xlsx"'}
    )

# ============= SERVICE/DEPARTMENT ROUTES =============
@api_router.get("/services")
async def get_services(current_user: dict = Depends(get_current_user)):
    """Get all services for the organization"""
    services = await db.services.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(100)
    return services

@api_router.post("/services")
async def create_service(service: ServiceCreate, current_user: dict = Depends(get_current_user)):
    """Create a new service/department"""
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can create services")
    
    # Check if service already exists
    existing = await db.services.find_one({"org_id": current_user['org_id'], "name": service.name})
    if existing:
        raise HTTPException(status_code=400, detail="Service with this name already exists")
    
    new_service = Service(**service.model_dump(), org_id=current_user['org_id'])
    await db.services.insert_one(new_service.model_dump())
    return {"message": "Service created successfully", "service": new_service}

@api_router.patch("/services/{service_id}")
async def update_service(service_id: str, service: ServiceCreate, current_user: dict = Depends(get_current_user)):
    """Update a service/department"""
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can update services")
    
    existing_service = await db.services.find_one({"id": service_id, "org_id": current_user['org_id']})
    if not existing_service:
        raise HTTPException(status_code=404, detail="Service not found")
    
    # Check if new name conflicts
    name_conflict = await db.services.find_one({
        "org_id": current_user['org_id'], 
        "name": service.name,
        "id": {"$ne": service_id}
    })
    if name_conflict:
        raise HTTPException(status_code=400, detail="Service with this name already exists")
    
    await db.services.update_one(
        {"id": service_id, "org_id": current_user['org_id']},
        {"$set": {"name": service.name}}
    )
    return {"message": "Service updated successfully"}

@api_router.delete("/services/{service_id}")
async def delete_service(service_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a service/department"""
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can delete services")
    
    result = await db.services.delete_one({"id": service_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Service not found")
    
    return {"message": "Service deleted successfully"}

# ============= CLIENT ONBOARDING ROUTES =============

@api_router.get("/client-onboarding", response_model=List[ClientOnboarding])
async def get_client_onboarding(current_user: dict = Depends(get_current_user)):
    onboardings = await db.client_onboarding.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    return onboardings

@api_router.post("/client-onboarding", response_model=ClientOnboarding)
async def create_client_onboarding(data: ClientOnboardingCreate, current_user: dict = Depends(get_current_user)):
    onboarding = ClientOnboarding(**data.model_dump(), org_id=current_user['org_id'])
    doc = onboarding.model_dump()
    await db.client_onboarding.insert_one(doc)
    return onboarding

@api_router.patch("/client-onboarding/{onboarding_id}")
async def update_client_onboarding(onboarding_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    onboarding = await db.client_onboarding.find_one({"id": onboarding_id, "org_id": current_user['org_id']})
    if not onboarding:
        raise HTTPException(status_code=404, detail="Onboarding not found in your organization")
    
    await db.client_onboarding.update_one({"id": onboarding_id, "org_id": current_user['org_id']}, {"$set": update_data})
    return {"message": "Onboarding updated successfully"}

@api_router.delete("/client-onboarding/{onboarding_id}")
async def delete_client_onboarding(onboarding_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.client_onboarding.delete_one({"id": onboarding_id, "org_id": current_user['org_id']})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Onboarding not found")
    return {"message": "Onboarding deleted successfully"}

# ============= CONSUMABLES ROUTES =============

@api_router.get("/stock-availability", response_model=List[StockAvailability])
async def get_stock_availability(current_user: dict = Depends(get_current_user)):
    stocks = await db.stock_availability.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    return stocks

@api_router.get("/stock-transactions", response_model=List[StockTransaction])
async def get_stock_transactions(current_user: dict = Depends(get_current_user)):
    transactions = await db.stock_transactions.find({"org_id": current_user['org_id']}, {"_id": 0}).to_list(1000)
    # Sort by date descending
    transactions.sort(key=lambda x: x.get('date', ''), reverse=True)
    return transactions

@api_router.post("/stock-in")
async def stock_in(data: StockInCreate, current_user: dict = Depends(get_current_user)):
    # Create transaction
    transaction = StockTransaction(
        org_id=current_user['org_id'],
        type='Stock In',
        product_name=data.product_name,
        vendor_name_or_issued_to=data.vendor_name,
        invoice_number=data.invoice_number,
        email=data.email,
        date=data.date,
        quantity=data.quantity,
        price=data.price
    )
    await db.stock_transactions.insert_one(transaction.model_dump())
    
    # Update stock availability
    existing_stock = await db.stock_availability.find_one({
        "product_name": data.product_name,
        "org_id": current_user['org_id']
    })
    
    if existing_stock:
        # Update quantity
        new_quantity = existing_stock['stock_available'] + data.quantity
        await db.stock_availability.update_one(
            {"id": existing_stock['id']},
            {"$set": {"stock_available": new_quantity, "vendor_name": data.vendor_name}}
        )
    else:
        # Create new stock
        stock = StockAvailability(
            org_id=current_user['org_id'],
            product_name=data.product_name,
            vendor_name=data.vendor_name,
            stock_available=data.quantity
        )
        await db.stock_availability.insert_one(stock.model_dump())
    
    return {"message": "Stock In recorded successfully"}

@api_router.post("/stock-out")
async def stock_out(data: StockOutCreate, current_user: dict = Depends(get_current_user)):
    # Check if stock exists and has enough quantity
    existing_stock = await db.stock_availability.find_one({
        "product_name": data.product_name,
        "org_id": current_user['org_id']
    })
    
    if not existing_stock:
        raise HTTPException(status_code=404, detail="Product not found in stock")
    
    if existing_stock['stock_available'] < data.quantity:
        raise HTTPException(status_code=400, detail=f"Insufficient stock. Available: {existing_stock['stock_available']}")
    
    # Create transaction
    transaction = StockTransaction(
        org_id=current_user['org_id'],
        type='Stock Out',
        product_name=data.product_name,
        vendor_name_or_issued_to=data.issued_to,
        email=data.email,
        date=data.date,
        quantity=data.quantity
    )
    await db.stock_transactions.insert_one(transaction.model_dump())
    
    # Update stock availability
    new_quantity = existing_stock['stock_available'] - data.quantity
    await db.stock_availability.update_one(
        {"id": existing_stock['id']},
        {"$set": {"stock_available": new_quantity}}
    )
    
    return {"message": "Stock Out recorded successfully"}

@api_router.patch("/stock-availability/{stock_id}")
async def update_stock_notes(stock_id: str, update_data: dict, current_user: dict = Depends(get_current_user)):
    stock = await db.stock_availability.find_one({"id": stock_id, "org_id": current_user['org_id']})
    if not stock:
        raise HTTPException(status_code=404, detail="Stock not found")
    
    await db.stock_availability.update_one({"id": stock_id}, {"$set": update_data})
    return {"message": "Stock updated successfully"}

@api_router.get("/stock-products")
async def get_stock_products(current_user: dict = Depends(get_current_user)):
    """Get list of all product names for dropdown"""
    stocks = await db.stock_availability.find(
        {"org_id": current_user['org_id']}, 
        {"_id": 0, "product_name": 1, "stock_available": 1}
    ).to_list(1000)
    return stocks

app.include_router(api_router)

# Mount static files for uploads
uploads_dir = ROOT_DIR / "uploads"
uploads_dir.mkdir(parents=True, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=uploads_dir), name="uploads")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ============= ADMIN UTILITIES =============
@api_router.post("/admin/clear-org-data")
async def clear_org_data(current_user: dict = Depends(get_current_user)):
    """Clear all data for the organization (Admin only)"""
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can clear organization data")
    
    org_id = current_user['org_id']
    
    # Delete all data for this organization
    await db.users.delete_many({"org_id": org_id})
    await db.clients.delete_many({"org_id": org_id})
    await db.contractors.delete_many({"org_id": org_id})
    await db.employees.delete_many({"org_id": org_id})
    await db.assets.delete_many({"org_id": org_id})
    await db.client_onboarding.delete_many({"org_id": org_id})
    await db.stock_availability.delete_many({"org_id": org_id})
    await db.stock_transactions.delete_many({"org_id": org_id})
    await db.services.delete_many({"org_id": org_id})
    
    return {"message": "All organization data cleared successfully"}

@api_router.post("/admin/initialize-services")
async def initialize_services(current_user: dict = Depends(get_current_user)):
    """Initialize default services for the organization (Admin only)"""
    if current_user['role'] != 'Admin':
        raise HTTPException(status_code=403, detail="Only Admin can initialize services")
    
    org_id = current_user['org_id']
    
    # Check if services already exist
    existing_count = await db.services.count_documents({"org_id": org_id})
    if existing_count > 0:
        raise HTTPException(status_code=400, detail="Services already exist. Delete them first if you want to reinitialize.")
    
    # Create default services
    default_services = [
        "PPC", "SEO", "Content", "Backlink", "Business Development", "Others"
    ]
    
    for service_name in default_services:
        service = Service(name=service_name, org_id=org_id)
        await db.services.insert_one(service.model_dump())
    
    return {"message": f"Initialized {len(default_services)} default services"}

@app.on_event("startup")
async def startup():
    logger.info("Application started successfully")
    # No seed data - fresh start

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
